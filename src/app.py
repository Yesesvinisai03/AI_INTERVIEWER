# ==================================================
# 🔒 DISABLE CHROMA TELEMETRY (MUST BE FIRST)
# ==================================================
import os
os.environ["CHROMA_TELEMETRY_ENABLED"] = "false"
os.environ["ANONYMIZED_TELEMETRY"] = "false"

# ==================================================
# 🔥 LOAD .env FIRST
# ==================================================
from dotenv import load_dotenv
load_dotenv()

# ==================================================
# Standard imports
# ==================================================
import io
import uuid
import numpy as np
import streamlit as st
import matplotlib.pyplot as plt
from PIL import Image
from io import BytesIO
from interviewer.whiteboard import analyze_whiteboard
from interviewer.evaluation import generate_downloadable_report


from docx import Document

# ==================================================
# App imports
# ==================================================
from interviewer.state import InterviewState, CandidateInfo, Turn
from interviewer.graph import build_graph, PERSONALITY_DIFFICULTY_BOUNDS
from interviewer.resume_parser import extract_resume_text
from interviewer.rag import index_resume_to_chroma
from interviewer.evaluation import evaluate_interview
from interviewer.learning_plan import build_learning_plan

GRAPH = build_graph()

st.set_page_config(page_title="AI Interviewer", layout="wide")

# ==================================================
# Helpers
# ==================================================
def normalize(result) -> InterviewState:
    if isinstance(result, InterviewState):
        return result
    return InterviewState(**dict(result))

# Updates session state with the latest InterviewState
def run_graph_once():
    new_state = GRAPH.invoke(st.session_state.state)
    st.session_state.state = normalize(new_state)

# Retrieve the most recent interviewer question
# Used for whiteboard context and follow-up analysis

def get_current_question_text(state) -> str:
    """
    Returns the latest interviewer (agent) question from history.
    """
    for t in reversed(state.history):
        if t.speaker == "agent":
            return t.text
    return ""


# ---------- SMALL PIE CHART ----------
def render_pie(title: str, data: dict):
    if not data:
        return

    labels = list(data.keys())
    values = list(data.values())

    fig, ax = plt.subplots(figsize=(2.6, 2.6))
    wedges, _ = ax.pie(
        values,
        startangle=90,
        wedgeprops=dict(width=0.45),
    )
    ax.set_title(title, fontsize=10)

    ax.legend(
        wedges,
        labels,
        loc="center left",
        bbox_to_anchor=(1, 0.5),
        fontsize=8
    )

    # ✅ Correct indentation + new API
    st.pyplot(fig, width="content")


# Generate final interview report as a DOCX file
# Includes scores, feedback, learning plan, and evaluation
# ---------- DOCX REPORT ----------
def generate_report(candidate, feedback, plan) -> BytesIO:
    doc = Document()
    doc.add_heading("AI Interview Report", level=1)

    doc.add_paragraph(f"Name: {candidate.name}")
    doc.add_paragraph(f"Role: {candidate.role}")
    doc.add_paragraph(f"Experience: {candidate.experience_level}")
    doc.add_paragraph("")

    doc.add_heading("Overall Score", level=2)
    doc.add_paragraph(f"{feedback.get('score', 0)} / 10")

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(feedback.get("summary", ""))

    doc.add_heading("Strengths", level=2)
    for s in feedback.get("strengths", []):
        doc.add_paragraph(f"• {s}")

    doc.add_heading("Improvements", level=2)
    for i in feedback.get("improvements", []):
        doc.add_paragraph(f"• {i}")

    doc.add_heading("Learning Roadmap", level=2)
    for step in plan.get("steps", []):
        doc.add_paragraph(f"→ {step}")

    doc.add_heading("Recommended Resources", level=2)
    for r in plan.get("resources", []):
        doc.add_paragraph(f"• {r}")

    # ================= FINAL EVALUATION =================
    final_eval_text = generate_downloadable_report(st.session_state.state)

    doc.add_heading("Final Evaluation", level=2)

    for para in final_eval_text.split("\n\n"):
        doc.add_paragraph(para)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ==================================================
# Session init
# ==================================================
if "state" not in st.session_state:
    st.session_state.state = None
if "show_whiteboard" not in st.session_state:
    st.session_state.show_whiteboard = False

# Collect candidate details and resume before starting interview
# Performs live validation and enables Start button only when valid
# START SCREEN
if st.session_state.state is None:
    st.title("🧠 AI Interviewer (RAG + Whiteboard)")

    # ---------- INPUTS (OUTSIDE FORM FOR LIVE VALIDATION) ----------
    name = st.text_input("Name *")
    role = st.text_input("Role *", placeholder="e.g. Software Engineer")

    exp = st.selectbox(
        "Experience Level *",
        ["", "Entry", "Mid", "Senior"]
    )

    personality = st.selectbox(
        "Choose Interviewer Personality *",
        ["", "Friendly Coach", "Strict FAANG", "Startup Hiring Manager", "Professor"]
    )

    resume = st.file_uploader(
        "Upload Resume * (PDF / DOCX / TXT)",
        type=["pdf", "docx", "txt"]
    )

    # ---------- LIVE VALIDATION ----------
    errors = []

    if not name.strip():
        errors.append("Name is required")
    if not role.strip():
        errors.append("Role is required")
    if not exp:
        errors.append("Experience level is required")
    if not personality:
        errors.append("Interviewer personality is required")
    if resume is None:
        errors.append("Resume is required")

    # ---------- SHOW ERRORS (ONLY WHEN INVALID) ----------
    if errors:
        for e in errors:
            st.warning(e)

    # ---------- START BUTTON ----------
    start = st.button(
        "Start Interview",
        disabled=bool(errors)
    )

    # ---------- ON START ----------
    if start and not errors:
        resume_text = extract_resume_text(resume)
        user_id = str(uuid.uuid4())[:8]

        index_resume_to_chroma(
            user_id=user_id,
            resume_text=resume_text
        )

        st.session_state.state = InterviewState(
            candidate=CandidateInfo(
                name=name.strip(),
                role=role.strip(),
                experience_level=exp,
                resume_text=resume_text,
                user_id=user_id,
                personality=personality,
            )
        )

        run_graph_once()
        st.rerun()

    st.stop()


# ==================================================
# MAIN INTERVIEW UI
# ==================================================
state: InterviewState = st.session_state.state

st.title("🧠 AI Interviewer")
st.caption(f"🎭 Interviewer: **{state.candidate.personality}**")


# ==================================================
# 🎚️ DIFFICULTY INDICATOR (TECH ROUND ONLY)
# ==================================================
if state.current_round == "technical" and state.phase in {"questioning", "closing"}:
    personality = state.candidate.personality
    min_d, max_d = PERSONALITY_DIFFICULTY_BOUNDS.get(personality, (1, 5))

    # simple color based on level
    lvl = int(getattr(state, "difficulty_level", min_d) or min_d)
    color = "green" if lvl <= (min_d + 1) else "orange" if lvl < max_d else "red"

    st.markdown(
        f"""
        <div style="padding:10px;border-radius:10px;border:1px solid #ddd; margin-bottom:10px;">
            🧠 <b>Difficulty Level:</b>
            <span style="color:{color}; font-weight:700;">{lvl} / {max_d}</span><br>
            🎭 <b>Personality:</b> {personality}<br>
            🧩 <b>Round:</b> Technical
        </div>
        """,
        unsafe_allow_html=True
    )

# ---------------- CHAT HISTORY ----------------
for t in state.history:
    with st.chat_message("assistant" if t.speaker == "agent" else "user"):
        st.markdown(t.text)

# Capture candidate response and send it to interview agent
# ==================================================
# INPUT
# ==================================================
if state.phase in {"questioning", "closing"}:
    msg = st.chat_input("Type your answer (you may ask a final question)…")

    if msg:
        state.history.append(Turn(speaker="candidate", text=msg))
        state.last_user_answer = msg
        state.pending_user_input = True
        st.session_state.state = state
        run_graph_once()
        st.rerun()

# Allow whiteboard access only during technical round
# ---- OPEN BUTTON (TECHNICAL ROUND ONLY) ----
if not st.session_state.show_whiteboard:

    if state.current_round != "technical":
        st.info("🧾 Whiteboard is available only during the Technical round.")
    else:
        if st.button("🧾 Open Whiteboard"):
            st.session_state.show_whiteboard = True
            st.rerun()

# ---- WHITEBOARD UI ----
if st.session_state.show_whiteboard:

    # 🔒 HARD BLOCK — ONLY TECHNICAL ROUND
    if state.current_round != "technical":
        st.warning("❌ Whiteboard is only allowed in the Technical round.")
        st.session_state.show_whiteboard = False
        st.rerun()

    st.divider()
    st.subheader("🧾 Whiteboard (Flowcharts / Diagrams)")


    # ✅ Real whiteboard
    st.components.v1.iframe(
        src="https://excalidraw.com",
        height=600,
    )

    st.info(
        "👉 Draw your diagram using boxes/arrows → Export as PNG → Upload below"
    )

    uploaded_diagram = st.file_uploader(
        "📤 Upload exported diagram (PNG)",
        type=["png"]
    )

    col1, col2 = st.columns(2)

    with col1:
        if st.button("🧠 Analyze Whiteboard"):
            if uploaded_diagram is not None:
                question_text = get_current_question_text(state)
                img = Image.open(uploaded_diagram)
                img_arr = np.array(img)

                result = analyze_whiteboard(
                    question=question_text,
                    candidate_text=state.last_user_answer or "",
                    canvas_image=img_arr,
                )

                # 1️⃣ Show feedback immediately
                if isinstance(result, dict) and "feedback" in result:
                    state.history.append(
                        Turn(
                            speaker="agent",
                            text=f"📝 **Whiteboard Feedback**\n\n{result['feedback']}"
                        )
                    )

                # 2️⃣ Ask follow-up question immediately
                if isinstance(result, dict) and "question" in result:
                    state.history.append(
                        Turn(
                            speaker="agent",
                            text=result["question"]
                        )
                    )

                # 3️⃣ Save state
                st.session_state.state = state

                st.success("Whiteboard analyzed using vision model")

                # 4️⃣ Close whiteboard & rerun
                st.session_state.show_whiteboard = False
                st.rerun()

            else:
                st.warning("Please upload the exported PNG from the whiteboard.")

    with col2:
        if st.button("❌ Close Whiteboard"):
            st.session_state.show_whiteboard = False
            st.rerun()



# ==================================================
# POST INTERVIEW
# ==================================================
if state.interview_done:
    st.divider()
    st.success("🎉 Interview Finished")

    if st.button("📋 Generate Feedback"):
        state.feedback = evaluate_interview(state.records)

        # ✅ ALWAYS initialize learning_plan
        state.learning_plan = build_learning_plan(
            role=state.candidate.role,
            evaluation=state.feedback,
            personality=state.candidate.personality,
        )

        st.session_state.state = state
        st.rerun()


    if state.feedback:
        st.subheader("📊 Overall Score")
        st.metric("Score", f"{state.feedback['score']}/10")
        

        # ===============================
        # 🧑‍💼 HR PERFORMANCE (⬅ ADD HERE)
        # ===============================
        if "hr_performance" in state.feedback:
            st.subheader("🧑‍💼 HR Performance")

            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric(
                    "Communication",
                    state.feedback["hr_performance"]["communication"]
                )
            with col2:
                st.metric(
                    "Clarity",
                    state.feedback["hr_performance"]["clarity"]
                )
            with col3:
                st.metric(
                    "Confidence",
                    state.feedback["hr_performance"]["confidence"]
                )

        # -------------------------------
        # Existing breakdowns
        # -------------------------------
        st.subheader("📈 Skill Breakdown")
        render_pie("Skill Dimensions", state.feedback.get("dimensions", {}))

        st.subheader("📚 Topic Performance")
        topic_scores = {
            k: v["good"] for k, v in state.feedback.get("topic_breakdown", {}).items()
        }
        render_pie("Topics", topic_scores)

        st.subheader("✅ Strengths")
        for s in state.feedback.get("strengths", []):
            st.success(s)

        st.subheader("⚠️ Improvements")
        for i in state.feedback.get("improvements", []):
            st.warning(i)

        st.subheader("📝 Summary")
        st.info(state.feedback.get("summary", ""))

        
        total_hints = state.feedback.get("total_hints", 0)

        if total_hints > 0:
            st.warning(f"💡 Hints used during interview: {total_hints}")
        else:
            st.success("💪 No hints used — excellent confidence!")

    # ✅ SAFE GUARD — learning_plan may be None
    if state.learning_plan:

        if "overall_focus" in state.learning_plan:
            st.info(state.learning_plan["overall_focus"])

        st.subheader("🗺️ Learning Roadmap")
        for step in state.learning_plan.get("steps", []):
            st.write("➡️", step)

        st.subheader("📚 Recommended Resources")

        for r in state.learning_plan.get("resources", []):
            # Heading line (no colon)
            if ":" not in r and len(r.split()) <= 3:
                st.markdown(f"**{r}**")
            # Resource line (clickable link)
            else:
                st.markdown(r)



        report = generate_report(state.candidate, state.feedback, state.learning_plan)
        st.download_button(
            "⬇️ Download Interview Report (DOCX)",
            data=report,
            file_name="interview_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

